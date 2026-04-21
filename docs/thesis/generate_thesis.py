#!/usr/bin/env python3
"""
Generator pracy licencjackiej: "Projekt i implementacja platformy do zarzadzania usluga VPN".
Generuje plik DOCX zgodny z wymaganiami edytorskimi KUL.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor
from docxcompose.composer import Composer

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parent.parent.parent
TITLE_PAGE = REPO_ROOT / "docs" / "strona_tytulowa_128197 (1).docx"
OUTPUT_FILE = REPO_ROOT / "docs" / "thesis" / "praca_licencjacka.docx"
BODY_FILE = REPO_ROOT / "docs" / "thesis" / "_praca_licencjacka_body.docx"

THESIS_FONT = "Times New Roman"
CODE_FONT = "Courier New"

STYLE_BODY = "ThesisBody"
STYLE_BODY_INDENT = "ThesisBodyIndent"
STYLE_LIST = "ThesisList"
STYLE_H1 = "ThesisHeading1"
STYLE_H2 = "ThesisHeading2"
STYLE_H3 = "ThesisHeading3"
STYLE_TOC_TITLE = "ThesisTocTitle"
STYLE_CODE_CAPTION = "ThesisCodeCaption"
STYLE_SOURCE_NOTE = "ThesisSourceNote"
STYLE_CODE = "ThesisCode"
STYLE_TABLE_CAPTION = "ThesisTableCaption"
STYLE_TABLE_TEXT = "ThesisTableText"
STYLE_FIGURE_CAPTION = "ThesisFigureCaption"
STYLE_PLACEHOLDER = "ThesisPlaceholder"
STYLE_BIB = "ThesisBibliography"

def read_code_excerpt(relative_path: str, ranges: list[tuple[int, int]]) -> str:
    """Read selected line ranges and prefix each line with its source line number."""
    full = REPO_ROOT / relative_path
    if not full.exists():
        return f"[BRAK PLIKU: {relative_path}]"

    lines = full.read_text(encoding="utf-8").splitlines()
    width = max(len(str(end)) for _, end in ranges)
    excerpt: list[str] = []

    for idx, (start, end) in enumerate(ranges):
        if idx:
            excerpt.append(f"{' ' * width} ...")
        for line_no in range(start, min(end, len(lines)) + 1):
            excerpt.append(f"{line_no:>{width}} {lines[line_no - 1]}")

    return "\n".join(excerpt)


def format_ranges(ranges: list[tuple[int, int]]) -> str:
    return ", ".join(
        f"{start}-{end}" if start != end else str(start) for start, end in ranges
    )


# ---------------------------------------------------------------------------
# Document formatting helpers
# ---------------------------------------------------------------------------

def configure_font(holder, name: str, size_pt: float, *, bold=None, italic=None):
    """Apply a font consistently to a style or run."""
    holder.font.name = name
    holder.font.size = Pt(size_pt)
    if bold is not None:
        holder.font.bold = bold
    if italic is not None:
        holder.font.italic = italic


def style_paragraph(style, *, alignment, line_spacing, before=0, after=0):
    pf = style.paragraph_format
    pf.alignment = alignment
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def set_style_outline_level(style, level: int):
    """Assign Word outline level so custom heading styles appear in TOC."""
    style_el = style.element
    ppr = style_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        style_el.append(ppr)
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def setup_styles(doc: Document):
    """Configure custom KUL-compliant body styles without touching title-page styles."""
    base = ensure_style(doc, STYLE_BODY)
    configure_font(base, THESIS_FONT, 12)
    base.font.color.rgb = RGBColor(0, 0, 0)
    style_paragraph(
        base,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5,
    )

    body_indent = ensure_style(doc, STYLE_BODY_INDENT)
    body_indent.base_style = base
    configure_font(body_indent, THESIS_FONT, 12)
    style_paragraph(
        body_indent,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5,
    )
    body_indent.paragraph_format.first_line_indent = Cm(1.25)

    list_style = ensure_style(doc, STYLE_LIST)
    list_style.base_style = base
    configure_font(list_style, THESIS_FONT, 12)
    style_paragraph(
        list_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
    )
    list_style.paragraph_format.left_indent = Cm(0.75)
    list_style.paragraph_format.first_line_indent = Cm(-0.5)

    h1 = ensure_style(doc, STYLE_H1)
    h1.base_style = base
    configure_font(h1, THESIS_FONT, 14, bold=True)
    style_paragraph(
        h1,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        before=12,
        after=10,
    )
    h1.paragraph_format.keep_with_next = True
    set_style_outline_level(h1, 0)

    h2 = ensure_style(doc, STYLE_H2)
    h2.base_style = base
    configure_font(h2, THESIS_FONT, 12, bold=True, italic=False)
    style_paragraph(
        h2,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        before=6,
        after=6,
    )
    h2.paragraph_format.keep_with_next = True
    set_style_outline_level(h2, 1)

    h3 = ensure_style(doc, STYLE_H3)
    h3.base_style = base
    configure_font(h3, THESIS_FONT, 12, bold=False, italic=True)
    style_paragraph(
        h3,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        before=6,
        after=6,
    )
    h3.paragraph_format.keep_with_next = True
    set_style_outline_level(h3, 2)

    toc_title = ensure_style(doc, STYLE_TOC_TITLE)
    toc_title.base_style = base
    configure_font(toc_title, THESIS_FONT, 14, bold=True)
    style_paragraph(
        toc_title,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        after=6,
    )

    code_caption = ensure_style(doc, STYLE_CODE_CAPTION)
    code_caption.base_style = base
    configure_font(code_caption, THESIS_FONT, 10, italic=True)
    style_paragraph(
        code_caption,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.0,
        before=6,
        after=2,
    )

    source_note = ensure_style(doc, STYLE_SOURCE_NOTE)
    source_note.base_style = base
    configure_font(source_note, THESIS_FONT, 9, italic=True)
    source_note.font.color.rgb = RGBColor(96, 96, 96)
    style_paragraph(
        source_note,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.0,
        after=2,
    )

    code_style = ensure_style(doc, STYLE_CODE)
    configure_font(code_style, CODE_FONT, 7.5)
    style_paragraph(
        code_style,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.0,
        after=6,
    )
    code_style.paragraph_format.left_indent = Cm(0.35)
    code_style.paragraph_format.right_indent = Cm(0.35)

    table_caption = ensure_style(doc, STYLE_TABLE_CAPTION)
    table_caption.base_style = base
    configure_font(table_caption, THESIS_FONT, 10, bold=True)
    style_paragraph(
        table_caption,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        before=8,
        after=4,
    )

    table_text = ensure_style(doc, STYLE_TABLE_TEXT)
    table_text.base_style = base
    configure_font(table_text, THESIS_FONT, 10)
    style_paragraph(
        table_text,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.0,
    )

    figure_caption = ensure_style(doc, STYLE_FIGURE_CAPTION)
    figure_caption.base_style = base
    configure_font(figure_caption, THESIS_FONT, 10, italic=True)
    style_paragraph(
        figure_caption,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        after=8,
    )

    placeholder = ensure_style(doc, STYLE_PLACEHOLDER)
    placeholder.base_style = base
    configure_font(placeholder, THESIS_FONT, 11)
    placeholder.font.color.rgb = RGBColor(128, 128, 128)
    style_paragraph(
        placeholder,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        before=12,
        after=4,
    )

    bibliography = ensure_style(doc, STYLE_BIB)
    bibliography.base_style = base
    configure_font(bibliography, THESIS_FONT, 12)
    style_paragraph(
        bibliography,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5,
        after=4,
    )
    bibliography.paragraph_format.first_line_indent = Cm(-1.25)
    bibliography.paragraph_format.left_indent = Cm(1.25)


def setup_page(doc: Document, section=None):
    """Set A4 page with KUL margins."""
    section = section or doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_page_numbers(doc: Document, section=None):
    """Add page numbers to footer."""
    section = section or doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_xml = (
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:instr="PAGE \\* MERGEFORMAT">'
        '<w:r><w:t>1</w:t></w:r>'
        '</w:fldSimple>'
    )
    run._r.append(parse_xml(fld_xml))


def enable_update_fields_on_open(doc: Document):
    """Tell Word to refresh fields such as TOC/page refs when the document is opened."""
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def strip_toc_bookmarks(doc: Document):
    """Remove stale _Toc bookmarks copied from the title page template."""
    for paragraph in doc.paragraphs:
        for child in list(paragraph._p):
            tag = child.tag.rsplit("}", 1)[-1]
            if tag not in {"bookmarkStart", "bookmarkEnd"}:
                continue
            name = child.get(qn("w:name"), "")
            if not name or name.startswith("_Toc"):
                paragraph._p.remove(child)


def add_toc(doc: Document):
    """Insert a Table of Contents field."""
    p = doc.add_paragraph("Spis treści", style=STYLE_TOC_TITLE)
    configure_font(p.runs[0], THESIS_FONT, 14, bold=True)

    p2 = doc.add_paragraph(style=STYLE_BODY)
    fld_xml = (
        '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:instr="TOC \\o &quot;1-3&quot; \\h \\z \\u">'
        '<w:r><w:rPr><w:i/></w:rPr><w:t>[Zaktualizuj spis treści: Ctrl+A, F9]</w:t></w:r>'
        '</w:fldSimple>'
    )
    p2._p.append(parse_xml(fld_xml))
    doc.add_page_break()


def p(doc: Document, text: str):
    """Add a body paragraph or a list item."""
    style = STYLE_LIST if text.lstrip().startswith("- ") else STYLE_BODY
    para = doc.add_paragraph(text, style=style)
    return para


def p_indent(doc: Document, text: str):
    """Add a paragraph with first-line indent."""
    return doc.add_paragraph(text, style=STYLE_BODY_INDENT)


def add_heading(doc: Document, text: str, level: int):
    style = {1: STYLE_H1, 2: STYLE_H2, 3: STYLE_H3}[level]
    return doc.add_paragraph(text, style=style)


def wrap_code_block(code: str, width: int = 96) -> str:
    wrapped_lines: list[str] = []
    for line in code.splitlines():
        current = line
        while len(current) > width:
            wrapped_lines.append(current[:width])
            current = " " * 8 + current[width:]
        wrapped_lines.append(current)
    return "\n".join(wrapped_lines)


def add_code(
    doc: Document,
    code: str,
    caption: str,
    listing_num: str,
    *,
    source_path: str | None = None,
    line_ranges: list[tuple[int, int]] | None = None,
):
    """Add a code listing with caption."""
    cap = doc.add_paragraph(style=STYLE_CODE_CAPTION)
    cap.add_run(f"Listing {listing_num}: {caption}")

    if source_path and line_ranges:
        src = doc.add_paragraph(
            f"Źródło: {source_path}, lin. {format_ranges(line_ranges)}.",
            style=STYLE_SOURCE_NOTE,
        )
        src.paragraph_format.keep_with_next = True

    code_para = doc.add_paragraph(style=STYLE_CODE)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
    code_para._p.get_or_add_pPr().append(shading)
    run = code_para.add_run(wrap_code_block(code))
    configure_font(run, CODE_FONT, 7.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str, table_num: str):
    """Add a formatted table with caption."""
    cap = doc.add_paragraph(f"Tabela {table_num}: {caption}", style=STYLE_TABLE_CAPTION)
    configure_font(cap.runs[0], THESIS_FONT, 10, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.paragraphs[0].style = STYLE_TABLE_TEXT
        rr = cell.paragraphs[0].add_run(h)
        configure_font(rr, THESIS_FONT, 10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            cell.paragraphs[0].style = STYLE_TABLE_TEXT
            rr = cell.paragraphs[0].add_run(val)
            configure_font(rr, THESIS_FONT, 10)

    doc.add_paragraph("", style=STYLE_BODY)


def add_figure_placeholder(doc: Document, caption: str, fig_num: str):
    """Add a placeholder for a figure or screenshot."""
    para = doc.add_paragraph(style=STYLE_PLACEHOLDER)

    # Placeholder box
    border_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="8" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="8" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="8" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="8" w:color="999999"/>'
        '</w:pBdr>'
    )
    para._p.get_or_add_pPr().append(parse_xml(border_xml))

    run = para.add_run("\n\n[PLACEHOLDER: Wstaw zrzut ekranu lub diagram]\n\n")
    configure_font(run, THESIS_FONT, 11)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Caption below
    cap = doc.add_paragraph(f"Rysunek {fig_num}: {caption}", style=STYLE_FIGURE_CAPTION)
    configure_font(cap.runs[0], THESIS_FONT, 10, italic=True)


# ===========================================================================
# CHAPTERS
# ===========================================================================

def chapter_1(doc: Document):
    """Rozdział 1: Wstęp"""
    add_heading(doc, "1. Wstęp", 1)

    p(doc,
      "Współczesny Internet, mimo że stanowi fundamentalne narzędzie komunikacji, pracy "
      "i rozrywki, nie zapewnia domyślnie prywatności swoim użytkownikom. Dane przesyłane "
      "przez sieć mogą być przechwytywane, analizowane i profilowane zarówno przez dostawców "
      "usług internetowych, jak i przez podmioty trzecie. W odpowiedzi na te zagrożenia "
      "coraz większą popularność zyskują wirtualne sieci prywatne (VPN), które tworzą "
      "szyfrowane tunele między urządzeniem użytkownika a serwerem pośredniczącym, "
      "ukrywając tym samym zawartość i metadane ruchu sieciowego.")

    p_indent(doc,
      "Rynek usług VPN rozwija się dynamicznie. Wzrost ten napędzany jest przez "
      "regulacje takie jak RODO (GDPR) w Europie, rosnącą świadomość zagrożeń "
      "cybernetycznych oraz potrzebę omijania blokad geograficznych. Jednocześnie "
      "coraz więcej usług przechodzi na model SaaS (Software as a Service), który "
      "umożliwia elastyczne zarządzanie subskrypcjami i skalowanie infrastruktury.")

    p_indent(doc,
      "Mimo dużej liczby komercyjnych dostawców VPN, większość z nich oferuje zamknięte, "
      "nieprzejrzyste rozwiązania, w których użytkownik nie ma możliwości weryfikacji, "
      "jakie dane są rzeczywiście zbierane i przetwarzane. Ponadto brakuje otwartych, "
      "kompletnych implementacji platformy SaaS do zarządzania infrastrukturą VPN, "
      "które mogłyby służyć zarówno jako rozwiązanie produkcyjne, jak i materiał "
      "edukacyjny.")

    add_heading(doc, "1.1. Cel pracy", 2)
    p(doc,
      "Celem niniejszej pracy jest zaprojektowanie i implementacja kompletnej platformy "
      "typu SaaS do zarządzania usługą VPN, nazwanej vpnVPN. Platforma ma umożliwiać "
      "użytkownikom końcowym zakup subskrypcji, zarządzanie urządzeniami i nawiązywanie "
      "szyfrowanych połączeń VPN, a operatorowi — centralne zarządzanie flotą serwerów "
      "VPN rozproszonych geograficznie na zróżnicowanej infrastrukturze.")

    p_indent(doc,
      "Cele szczegółowe pracy obejmują: (1) analizę i porównanie protokołów VPN pod kątem "
      "wydajności, bezpieczeństwa i kompatybilności; (2) zaprojektowanie architektury "
      "wielokomponentowego systemu z podziałem na warstwę prezentacji, sterowania "
      "i danych; (3) implementację czterech głównych komponentów systemu; (4) przeprowadzenie "
      "testów integracyjnych i wdrożenie na środowisko produkcyjne.")

    add_heading(doc, "1.2. Zakres pracy", 2)
    p(doc,
      "Platforma vpnVPN składa się z czterech głównych komponentów:")

    p(doc,
      "- Aplikacja webowa — interfejs oparty na Next.js 15, odpowiedzialny za "
      "rejestrację i logowanie użytkowników, obsługę płatności (Stripe), panel użytkownika "
      "z zarządzaniem urządzeniami i konfiguracjami VPN oraz panel administracyjny.")

    p(doc,
      "- Warstwa sterująca — serwis HTTP zbudowany na Bun/Fastify z bazą danych "
      "PostgreSQL, przechowująca stan systemu i udostępniająca API REST do rejestracji "
      "węzłów, synchronizacji peerów i zarządzania tokenami.")

    p(doc,
      "- Serwer VPN — agent napisany w języku Rust, uruchamiany jako kontener Docker "
      "na maszynach wirtualnych, obsługujący trzy protokoły VPN: WireGuard, OpenVPN "
      "i IKEv2/IPsec.")

    p(doc,
      "- Aplikacja desktopowa — klient Tauri z uprzywilejowanym daemonem Rust, "
      "obsługująca WireGuard, OpenVPN i IKEv2 na systemach macOS, Linux i Windows.")

    p_indent(doc,
      "W pracy szczególny nacisk położono na prywatność użytkowników (brak logowania "
      "treści ruchu), bezpieczeństwo komunikacji (nowoczesne algorytmy "
      "kryptograficzne) oraz przejrzystą architekturę umożliwiającą weryfikację przepływu "
      "danych.")

    add_heading(doc, "1.3. Struktura pracy", 2)
    p(doc,
      "Praca składa się z sześciu rozdziałów. Rozdział 1 wprowadza w tematykę i definiuje "
      "cele pracy. Rozdział 2 prezentuje podstawy teoretyczne — protokoły VPN, technologie "
      "sieciowe, model SaaS oraz wymagania niefunkcjonalne. Rozdział 3 zawiera analizę "
      "wymagań i projekt systemu, w tym architekturę, model danych i porównanie technologii. "
      "Rozdział 4 opisuje implementację poszczególnych komponentów z fragmentami kodu. "
      "Rozdział 5 przedstawia integrację, testowanie i wdrożenie. Rozdział 6 zawiera "
      "podsumowanie i kierunki dalszego rozwoju.")


def chapter_2(doc: Document):
    """Rozdział 2: Podstawy teoretyczne"""
    add_heading(doc, "2. Podstawy teoretyczne", 1)

    # 2.1 Protokoły VPN
    add_heading(doc, "2.1. Przegląd protokołów VPN", 2)
    p(doc,
      "Wirtualna sieć prywatna (VPN, Virtual Private Network) to technologia pozwalająca "
      "na utworzenie szyfrowanego tunelu pomiędzy urządzeniem klienckim a serwerem VPN. "
      "Ruch sieciowy przechodzący przez tunel jest zabezpieczony przed podsłuchem i "
      "manipulacją, a adres IP użytkownika jest zastępowany adresem serwera VPN. "
      "W niniejszym podrozdziale omówiono trzy protokoły VPN wykorzystywane w systemie "
      "vpnVPN.")

    add_heading(doc, "2.1.1. WireGuard", 3)
    p(doc,
      "WireGuard to nowoczesny protokół VPN opracowany przez Jasona Donenfelda, "
      "którego pierwsza stabilna wersja została włączona do jądra systemu Linux w marcu "
      "2020 roku (wersja 5.6). Protokół wyróżnia się minimalizmem — jego implementacja "
      "jądrowa liczy około 4000 linii kodu, co jest kilkukrotnie mniej niż w przypadku "
      "OpenVPN czy IPsec.")

    p_indent(doc,
      "WireGuard wykorzystuje protokół Noise IK jako ramy dla uzgadniania kluczy, "
      "a konkretnie algorytmy Curve25519 (wymiana kluczy Diffiego-Hellmana na krzywych "
      "eliptycznych), ChaCha20-Poly1305 (szyfrowanie strumieniowe z uwierzytelnianiem "
      "AEAD), BLAKE2s (funkcja haszująca) oraz SipHash24 (tablica haszująca). Taki dobór "
      "kryptoprimitywy eliminuje problem negocjacji parametrów — wszystkie połączenia "
      "WireGuard korzystają z tego samego zestawu algorytmów, co znacząco redukuje "
      "powierzchnię ataku.")

    p_indent(doc,
      "Z punktu widzenia wydajności WireGuard działa w warstwie 3 modelu OSI (IP), "
      "wykorzystuje wyłącznie protokół UDP i interfejsy TUN. Mechanizm routingu "
      "oparty jest na Cryptokey Routing Table — każdemu peerowi przypisana jest para "
      "klucz publiczny plus dozwolone adresy IP. Protokół nie wymaga zestawiania "
      "połączenia w tradycyjnym sensie — wymiana pakietów inicjujących (handshake) "
      "uruchamiana jest automatycznie przy pierwszym pakiecie danych, a materiał "
      "sesyjny jest okresowo odświeżany przez sam protokół.")

    p_indent(doc,
      "Istotną cechą WireGuard jest podejście do wersjonowania kryptografii. W przeciwieństwie "
      "do protokołów takich jak TLS czy IPsec, WireGuard nie oferuje negocjacji algorytmów — "
      "używa wyłącznie jednego, starannie dobranego zestawu kryptoprimitywy. Gdy algorytmy "
      "staną się przestarzałe, publikowana jest nowa wersja protokołu z nowym zestawem. "
      "Takie podejście eliminuje ryzyko ataków typu downgrade, w których napastnik wymusza "
      "użycie słabszego algorytmu.")

    p_indent(doc,
      "WireGuard implementuje mechanizm roamingu — jeśli peer zmieni adres IP (np. "
      "przełączenie z WiFi na LTE), wystarczy, że wyśle zaszyfrowany pakiet z nowego adresu, "
      "a serwer automatycznie zaktualizuje endpoint peera. Mechanizm ten, w połączeniu "
      "z PersistentKeepalive (wysyłanie pakietów podtrzymujących co 25 sekund w konfiguracji "
      "vpnVPN), zapewnia stabilne połączenie nawet w warunkach zmiennej łączności sieciowej.")

    add_heading(doc, "2.1.2. OpenVPN", 3)
    p(doc,
      "OpenVPN jest dojrzałym, otwartym rozwiązaniem VPN działającym w przestrzeni "
      "użytkownika (userspace), rozwijany od 2001 roku. Protokół wykorzystuje bibliotekę "
      "OpenSSL do negocjacji sesji TLS i szyfrowania danych, co zapewnia szerokie wsparcie "
      "dla różnych algorytmów kryptograficznych.")

    p_indent(doc,
      "W konfiguracji zastosowanej w systemie vpnVPN OpenVPN używa szyfrowania "
      "AES-256-GCM z dopuszczeniem CHACHA20-POLY1305, protokołu transportowego UDP "
      "na porcie 1194 oraz mechanizmu tls-crypt. W kodzie minimalna wersja TLS została "
      "ustawiona na 1.2, a mechanizm tls-crypt chroni kanał sterujący dodatkowym kluczem "
      "symetrycznym. Uwierzytelnianie klientów odbywa się za pomocą "
      "nazwy użytkownika i hasła weryfikowanego po stronie serwera przy użyciu skryptu "
      "auth-user-pass-verify z solonym haszem SHA-256.")

    p_indent(doc,
      "Główną zaletą OpenVPN jest dojrzałość ekosystemu i szeroka kompatybilność "
      "z urządzeniami klienckimi. Wadą jest większa złożoność konfiguracji, wyższy "
      "narzut wydajnościowy (praca w przestrzeni użytkownika) oraz znacząco większa "
      "baza kodu w porównaniu z WireGuard.")

    p_indent(doc,
      "OpenVPN obsługuje dwa tryby działania: routowany (dev tun, warstwa 3) i mostkowany "
      "(dev tap, warstwa 2). W systemie vpnVPN wykorzystywany jest tryb routowany, "
      "ponieważ tunelowanie odbywa się na poziomie pakietów IP i nie wymaga "
      "przekazywania ramek Ethernet. Tryb routowany jest również bardziej wydajny "
      "ze względu na mniejszy narzut nagłówkowy.")

    p_indent(doc,
      "Mechanizm tls-crypt zastosowany w vpnVPN dodaje dodatkową warstwę ochrony "
      "poprzez szyfrowanie całego kanału sterującego TLS kluczem symetrycznym współdzielonym "
      "między serwerem a klientem. Dzięki temu: (1) napastnik nie może zidentyfikować "
      "ruchu OpenVPN na podstawie nagłówków TLS (utrudnienie ataków DPI), (2) ataki "
      "denial-of-service na kanał sterujący wymagają znajomości klucza tls-crypt, "
      "(3) certyfikat serwera nie jest ujawniany w trakcie handshake'u TLS.")

    add_heading(doc, "2.1.3. IKEv2/IPsec", 3)
    p(doc,
      "IKEv2 (Internet Key Exchange version 2, RFC 7296) to protokół służący do "
      "negocjacji parametrów bezpieczeństwa i wymiany kluczy w ramach pakietu IPsec. "
      "W systemie vpnVPN implementacja IKEv2 opiera się na strongSwan — otwartej "
      "implementacji IPsec dla systemów Linux, macOS i Windows.")

    p_indent(doc,
      "Kluczową zaletą IKEv2 jest natywne wsparcie w większości systemów operacyjnych "
      "(macOS, iOS, Windows, Android), co eliminuje konieczność instalowania dodatkowego "
      "oprogramowania klienckiego. Protokół obsługuje MOBIKE (Mobility and Multihoming "
      "Protocol, RFC 4555), co pozwala na płynne przełączanie między sieciami (np. WiFi "
      "i LTE) bez zrywania połączenia VPN.")

    p_indent(doc,
      "W konfiguracji vpnVPN warstwa IKE wykorzystuje propozycje "
      "aes256-sha256-modp2048 oraz aes256-sha256-ecp256, natomiast dla ruchu ESP "
      "wymuszono aes256gcm128. Serwer uwierzytelnia się certyfikatem ECDSA P-256 "
      "podpisanym przez wewnętrzny urząd certyfikacji, a podstawowy profil klienta "
      "korzysta z EAP-MSCHAPv2.")

    p_indent(doc,
      "Proces nawiązywania połączenia IKEv2 składa się z dwóch faz: IKE_SA_INIT "
      "(negocjacja parametrów kryptograficznych i wymiana kluczy Diffiego-Hellmana) "
      "oraz IKE_AUTH (uwierzytelnienie stron i ustanowienie pierwszego Child SA). "
      "W konfiguracji vpnVPN serwer uwierzytelnia się certyfikatem ECDSA P-256, "
      "a klient domyślnie za pomocą EAP-MSCHAPv2 (login i hasło). Kod generuje "
      "również alternatywny profil EAP-TLS. Sama konfiguracja strongSwan powstaje "
      "w pliku ipsec.conf, natomiast narzędzie swanctl wykorzystywane jest do "
      "odświeżania stanu i wczytywania sekretów.")

    p_indent(doc,
      "Wybór ECDSA P-256 zamiast RSA wynika z implementacji generatora certyfikatów "
      "w module PKI oraz z potrzeby utrzymania jednego, spójnego zestawu certyfikatów "
      "dla OpenVPN i IKEv2. W repozytorium nie ma osobnej, ręcznie utrzymywanej "
      "infrastruktury klucza publicznego dla strongSwan.")

    # Tabela porównawcza
    add_table(doc,
        ["Protokół", "Algorytmy kryptograficzne", "Główne zalety", "Główne wady", "Zastosowanie w vpnVPN"],
        [
            ["WireGuard", "ChaCha20-Poly1305, Curve25519, BLAKE2s",
             "Wysoka wydajność, minimalny kod (~4000 LOC), prosta konfiguracja",
             "Młodszy protokół, brak negocjacji parametrów",
             "Protokół podstawowy, najwyższa wydajność"],
            ["OpenVPN", "AES-256-GCM, CHACHA20-POLY1305, TLS >= 1.2, tls-crypt",
             "Dojrzały ekosystem (20+ lat), szeroka kompatybilność",
             "Złożona konfiguracja, większy narzut, większa powierzchnia ataku",
             "Kompatybilność z istniejącymi klientami"],
            ["IKEv2/IPsec", "IKE: AES-256/SHA-256/MODP2048 lub ECP256; ESP: AES-256-GCM; EAP-MSCHAPv2",
             "Natywne wsparcie systemów operacyjnych, MOBIKE",
             "Złożoność implementacji, zależność od IPsec",
             "Klienci bez dodatkowego oprogramowania"],
        ],
        "Porównanie protokołów VPN wykorzystywanych w systemie vpnVPN", "1")

    # 2.2 Technologie sieciowe
    add_heading(doc, "2.2. Technologie sieciowe dla VPN", 2)
    p(doc,
      "Działanie sieci VPN opiera się na kilku fundamentalnych mechanizmach sieciowych, "
      "które umożliwiają tunelowanie, routing i translację adresów.")

    add_heading(doc, "2.2.1. Interfejsy TUN/TAP", 3)
    p(doc,
      "Interfejsy TUN i TAP to wirtualne urządzenia sieciowe dostępne w jądrze systemu "
      "Linux (oraz w innych systemach uniksowych). Interfejs TUN (tunnel) operuje w "
      "warstwie 3 modelu OSI — przyjmuje i generuje pakiety IP. Interfejs TAP (network "
      "tap) operuje w warstwie 2 — przetwarza ramki Ethernet. W systemie vpnVPN "
      "wszystkie trzy protokoły VPN korzystają z interfejsów TUN (urządzenie /dev/net/tun), "
      "ponieważ tunelowanie odbywa się na poziomie pakietów IP.")

    add_heading(doc, "2.2.2. Routing i NAT", 3)
    p(doc,
      "Po zestawieniu interfejsu tunelowego konieczne jest skonfigurowanie routingu, "
      "aby pakiety od klientów VPN mogły dotrzeć do Internetu. W systemie vpnVPN "
      "serwer VPN pełni rolę bramy domyślnej dla klientów — ruch kierowany jest do "
      "interfejsu tunelowego (np. wg0 dla WireGuard), a następnie przesyłany do "
      "domyślnego interfejsu sieciowego hosta za pomocą mechanizmu NAT (Network Address "
      "Translation) w trybie maskaradowania (MASQUERADE).")

    p_indent(doc,
      "Konfiguracja wymaga: (1) włączenia przekazywania pakietów IPv4 i IPv6 "
      "(sysctl net.ipv4.ip_forward=1, net.ipv6.conf.all.forwarding=1); (2) dodania "
      "reguły iptables -t nat -A POSTROUTING -o <interfejs> -j MASQUERADE; (3) zezwolenia "
      "na ruch pomiędzy interfejsem VPN a interfejsem domyślnym w łańcuchu FORWARD.")

    add_heading(doc, "2.2.3. MTU i fragmentacja", 3)
    p(doc,
      "Tunelowanie wprowadza dodatkowy narzut nagłówkowy, co wymaga odpowiedniego "
      "ustawienia MTU (Maximum Transmission Unit). Standardowe MTU dla Ethernetu wynosi "
      "1500 bajtów. WireGuard dodaje 60 bajtów narzutu (nagłówek UDP + WireGuard), "
      "dlatego w systemie vpnVPN MTU interfejsu wg0 ustawione jest na 1420 bajtów. "
      "OpenVPN i IKEv2/IPsec mają podobne wymagania, choć dokładne wartości zależą "
      "od wybranego szyfrowania i trybu transportu.")

    add_heading(doc, "2.2.4. Adresacja IPv4 i IPv6", 3)
    p(doc,
      "System vpnVPN przydziela klientom adresy z prywatnych pul: 10.8.0.0/24 dla "
      "WireGuard, 10.9.0.0/24 dla OpenVPN oraz domyślnie 10.9.0.0/24 także dla IKEv2. "
      "Adresacja IPv6 fd42:42:42::/64 jest używana zarówno przez WireGuard, jak i "
      "domyślną konfigurację IKEv2. W aplikacji webowej adres klienta dobierany jest "
      "deterministycznie na podstawie identyfikatora urządzenia i protokołu.")

    # 2.3 Architektury
    add_heading(doc, "2.3. Architektury infrastrukturalne dla usług VPN", 2)
    p(doc,
      "Usługi VPN mogą być wdrażane w różnych topologiach, w zależności od wymagań "
      "dotyczących skalowalności, opóźnień i redundancji:")

    p(doc,
      "- Hub-and-spoke — centralna architektura, w której wszystkie węzły VPN "
      "komunikują się z centralnym serwerem sterującym. Taka topologia jest prosta "
      "w zarządzaniu, ale centralny punkt stanowi single point of failure.")

    p(doc,
      "- Multi-region — rozproszona architektura z węzłami VPN w wielu regionach "
      "geograficznych. Użytkownicy mogą wybierać serwer najbliższy ich lokalizacji, "
      "co minimalizuje opóźnienia. System vpnVPN stosuje tę architekturę z centralnym "
      "Control Plane i rozproszonymi węzłami VPN.")

    p(doc,
      "- Mesh — każdy węzeł komunikuje się bezpośrednio z innymi węzłami, co eliminuje "
      "single point of failure, ale znacząco zwiększa złożoność zarządzania.")

    p_indent(doc,
      "W systemie vpnVPN zastosowano architekturę hub-and-spoke z elementami multi-region: "
      "centralny Control Plane (wdrożony na Railway) zarządza flotą węzłów VPN "
      "rozproszonych na maszynach wirtualnych w różnych lokalizacjach. Węzły samodzielnie "
      "rejestrują się w Control Plane i cyklicznie synchronizują listę dozwolonych "
      "peerów.")

    # 2.4 Technologie implementacji
    add_heading(doc, "2.4. Technologie i biblioteki do implementacji serwera VPN", 2)
    p(doc,
      "Wybór języka programowania i bibliotek ma kluczowe znaczenie dla bezpieczeństwa "
      "i wydajności serwera VPN. Poniższa tabela zestawia technologie faktycznie "
      "wykorzystane w repozytorium oraz ich rolę w systemie.")

    add_table(doc,
        ["Warstwa", "Rozwiązanie użyte w repozytorium", "Rola", "Ograniczenia"],
        [
            ["Warstwa asynchroniczna", "tokio",
             "Uruchamianie zadań współbieżnych w serwerze VPN",
             "Większa złożoność niż kod sekwencyjny"],
            ["Warstwa HTTP", "axum + reqwest",
             "Admin API po stronie węzła i komunikacja z Control Plane",
             "Wymaga jawnego modelowania kontraktów JSON"],
            ["WireGuard", "binarka systemowa wg oraz menedżer interfejsów systemowych",
             "Generowanie kluczy, status i aktualizacja peerów przez syncconf",
             "Zależność od narzędzi i uprawnień systemowych"],
            ["OpenVPN / IKEv2", "openvpn, strongSwan, swanctl",
             "Obsługa protokołów zgodnych z klientami legacy i natywnymi",
             "Większa złożoność operacyjna niż w przypadku WireGuard"],
        ],
        "Kluczowe technologie użyte w serwerze VPN", "2")

    # 2.5 Model SaaS
    add_heading(doc, "2.5. Model SaaS i usługi subskrypcyjne", 2)
    p(doc,
      "SaaS (Software as a Service) to model dystrybucji oprogramowania, w którym "
      "aplikacja jest hostowana centralnie i udostępniana użytkownikom przez Internet, "
      "najczęściej na zasadach subskrypcji. W kontekście usług VPN model SaaS oferuje "
      "kilka kluczowych korzyści:")

    p(doc,
      "- Automatyczne aktualizacje — użytkownicy nie muszą ręcznie aktualizować "
      "oprogramowania serwerowego.")

    p(doc,
      "- Skalowalność — operator może dynamicznie dodawać i usuwać węzły VPN "
      "w odpowiedzi na zapotrzebowanie.")

    p(doc,
      "- Elastyczne modele cenowe — wielopoziomowe plany subskrypcyjne pozwalają "
      "dostosować ofertę do różnych segmentów klientów.")

    p_indent(doc,
      "System vpnVPN implementuje model SaaS z czterema planami: Free (bezpłatny, "
      "1 urządzenie, ograniczona przepustowość), Basic (10 USD/miesiąc, 1 urządzenie), "
      "Pro (30 USD/miesiąc, 5 urządzeń) i Enterprise (1000 USD/miesiąc, limit 999 "
      "urządzeń w bieżącej implementacji). Obsługa płatności realizowana jest przez Stripe z mechanizmem "
      "webhooków do automatycznej synchronizacji statusu subskrypcji.")

    p_indent(doc,
      "Architektura SaaS w vpnVPN opiera się na wzorcu event-driven: zdarzenia Stripe "
      "(checkout.session.completed, customer.subscription.updated, customer.subscription.deleted) "
      "są przetwarzane asynchronicznie przez webhook handler. Każde zdarzenie wyzwala "
      "odpowiednie akcje: aktywację subskrypcji i wysłanie e-maila powitalnego, "
      "aktualizację planu, lub anulowanie subskrypcji z natychmiastowym odwołaniem "
      "wszystkich peerów VPN użytkownika. Taki model zapewnia spójność stanu systemu "
      "nawet w przypadku awarii lub opóźnień w przetwarzaniu.")

    p_indent(doc,
      "Wielopoziomowa struktura planów (tiering) implementowana jest przez mapowanie "
      "identyfikatorów cen Stripe (priceId) na konfigurację planu (limit urządzeń, "
      "dostępne funkcje). Mapowanie przechowywane jest w pliku tiers.ts i jest "
      "wykorzystywane zarówno po stronie frontendu (wyświetlanie cennika), jak i "
      "backendu (weryfikacja limitów przy rejestracji urządzeń).")

    # 2.6 Wymagania niefunkcjonalne
    add_heading(doc, "2.6. Wymagania niefunkcjonalne systemów VPN", 2)
    p(doc,
      "Systemy VPN, ze względu na przetwarzanie wrażliwych danych sieciowych, muszą "
      "spełniać szereg wymagań niefunkcjonalnych:")

    p(doc,
      "Prywatność — brak logowania ruchu sieciowego (no-logging policy) jest fundamentalnym "
      "wymogiem. System powinien przechowywać jedynie dane niezbędne do zestawienia "
      "tunelu i rozliczenia usługi, przy czym dla protokołów opartych na loginie i haśle "
      "oznacza to również przechowywanie danych uwierzytelniających w bazie.")

    p(doc,
      "Bezpieczeństwo — połączenia produkcyjne powinny być zabezpieczone protokołem TLS, "
      "a materiał kryptograficzny (klucze prywatne, hasła i sekretne pliki konfiguracyjne) "
      "przechowywany z restrykcyjnymi uprawnieniami i usuwany z pamięci po użyciu tam, "
      "gdzie implementacja to przewiduje.")

    p(doc,
      "Skalowalność — architektura powinna umożliwiać dodawanie nowych węzłów VPN "
      "bez przerw w działaniu usługi i bez modyfikacji istniejących komponentów.")

    p(doc,
      "Dostępność — system powinien być odporny na awarie pojedynczych komponentów "
      "i umożliwiać automatyczne wykrywanie nieosiągalnych węzłów.")

    p_indent(doc,
      "W kontekście systemów VPN szczególne znaczenie ma zasada privacy by design, "
      "zgodna z duchem RODO (Rozporządzenie 2016/679). Oznacza to, że mechanizmy "
      "ochrony prywatności powinny być wbudowane w architekturę systemu od samego "
      "początku, a nie dodawane ex post. W praktyce przekłada się to na decyzje "
      "projektowe takie jak: konfiguracja protokołów VPN z wyłączonym logowaniem, "
      "przechowywanie zagregowanych metryk (bez adresów IP użytkowników w module metryk), "
      "automatyczne zerowanie materiału kryptograficznego z pamięci oraz "
      "minimalizacja danych przekazywanych między komponentami systemu.")

    p_indent(doc,
      "Aspekt skalowalności realizowany jest w vpnVPN przez architekturę typu pull: "
      "węzły VPN samodzielnie pobierają konfigurację z Control Plane, co eliminuje "
      "problem aktywnego wypychania konfiguracji do wszystkich węzłów. Dodanie nowego "
      "węzła wymaga jedynie uruchomienia kontenera "
      "Docker z odpowiednimi zmiennymi środowiskowymi — węzeł automatycznie rejestruje "
      "się i rozpoczyna obsługę klientów.")


def chapter_3(doc: Document):
    """Rozdział 3: Analiza wymagań i projekt systemu"""
    add_heading(doc, "3. Analiza wymagań i projekt systemu vpnVPN", 1)

    # 3.1 Analiza funkcjonalna
    add_heading(doc, "3.1. Analiza funkcjonalna", 2)
    p(doc,
      "System vpnVPN obsługuje dwie podstawowe role: użytkownika końcowego oraz "
      "administratora. Każda z ról ma zdefiniowany zestaw dostępnych funkcji.")

    p(doc, "Użytkownik końcowy może:")
    p(doc, "- zarejestrować się i zalogować (OAuth: GitHub, Google; magic link e-mail),")
    p(doc, "- wykupić subskrypcję w jednym z czterech planów (Free, Basic, Pro, Enterprise),")
    p(doc, "- zarządzać urządzeniami VPN (dodawanie, usuwanie, pobieranie konfiguracji),")
    p(doc, "- przeglądać listę dostępnych serwerów VPN z ich statusem i metrykami,")
    p(doc, "- zarządzać ustawieniami konta i preferencjami powiadomień.")

    p(doc, "Administrator może:")
    p(doc, "- monitorować flotę serwerów VPN (status, metryki, obciążenie),")
    p(doc, "- zarządzać tokenami rejestracyjnymi dla węzłów VPN,")
    p(doc, "- przeglądać katalog użytkowników z informacjami o subskrypcjach,")
    p(doc, "- analizować zagregowane statystyki (sesje, dystrybucja geograficzna, kondycja serwerów).")

    p_indent(doc,
      "Kluczowe przypadki użycia obejmują: (1) zakup lub zmianę planu subskrypcyjnego "
      "przez użytkownika, (2) rejestrację urządzenia i pobranie konfiguracji VPN, "
      "(3) potwierdzenie lub anulowanie połączenia po stronie aplikacji desktopowej, "
      "(4) samoobsługową rejestrację nowego węzła VPN oraz (5) operacje administracyjne "
      "na serwerach i tokenach rejestracyjnych.")

    # 3.2 Wymagania niefunkcjonalne
    add_heading(doc, "3.2. Wymagania niefunkcjonalne", 2)

    add_table(doc,
        ["Wymaganie", "Opis", "Realizacja w vpnVPN"],
        [
            ["Prywatność", "Brak logowania ruchu, minimalizacja metadanych",
             "Brak logowania treści ruchu; przechowywane są jednak dane niezbędne do zestawienia tuneli i subskrypcji"],
            ["Bezpieczeństwo", "Szyfrowanie end-to-end, ochrona materiału kryptograficznego",
             "Klucze prywatne z prawami 0600, certyfikaty ECDSA P-256, częściowe zerowanie pamięci (Zeroizing)"],
            ["Skalowalność", "Możliwość dodawania węzłów bez przestojów",
             "Samoobsługowa rejestracja węzłów, centralne zarządzanie peerami"],
            ["Dostępność", "Odporność na awarie komponentów",
             "Heartbeat co 2 minuty, automatyczna detekcja offline po 5 minutach"],
            ["Wydajność", "Niskie opóźnienia, wysoka przepustowość",
             "WireGuard jako protokół domyślny, osobne ścieżki dla protokołów zgodnościowych"],
        ],
        "Wymagania niefunkcjonalne systemu vpnVPN", "3")

    # 3.3 Architektura
    add_heading(doc, "3.3. Architektura systemu", 2)
    p(doc,
      "System vpnVPN składa się z czterech głównych komponentów komunikujących się "
      "przez REST API, tRPC oraz lokalne IPC. Poniższy diagram przedstawia architekturę "
      "wysokiego poziomu.")

    add_figure_placeholder(doc, "Architektura systemu vpnVPN — cztery główne komponenty", "1")

    p(doc,
      "Aplikacja webowa (apps/web) — aplikacja Next.js 15 wdrożona na platformie Vercel. "
      "Na potrzeby operacji administracyjnych i rejestracji urządzeń komunikuje się z warstwą "
      "sterującą przez API REST z uwierzytelnianiem kluczem API "
      "(nagłówek x-api-key). Odpowiada za interfejs użytkownika i administratora, "
      "uwierzytelnianie (NextAuth.js), obsługę płatności (Stripe) i wysyłkę e-maili "
      "(Resend).")

    p(doc,
      "Warstwa sterująca (services/control-plane) — serwis HTTP oparty na Bun/Fastify, "
      "wdrożony na platformie Railway. Stanowi centralny punkt zarządzania stanem systemu. "
      "Przechowuje informacje o serwerach VPN, peerach, tokenach i metrykach w bazie "
      "PostgreSQL. Udostępnia dwa rodzaje uwierzytelniania: Bearer token dla "
      "węzłów VPN oraz klucz API dla aplikacji webowej.")

    p(doc,
      "Serwer VPN (apps/vpn-server) — binarka Rust uruchamiana jako kontener Docker "
      "z uprawnieniami NET_ADMIN. Obsługuje trzy backendy VPN (WireGuard, OpenVPN, IKEv2), "
      "rejestruje się w Control Plane, cyklicznie synchronizuje listę peerów i raportuje "
      "metryki. Udostępnia pomocnicze API administracyjne z endpointami /health, /metrics, "
      "/status, /pubkey i /logs.")

    p(doc,
      "Aplikacja desktopowa (apps/desktop) — aplikacja Tauri z podziałem na nieuprzywilejowany "
      "interfejs graficzny (React + Vite) i uprzywilejowany daemon Rust komunikujący się "
      "przez gniazdo Unix lub nazwany kanał, z komunikacją w formacie JSON-RPC. "
      "Daemon zarządza połączeniami VPN na poziomie "
      "systemowym.")

    # 3.4 Model danych
    add_heading(doc, "3.4. Model danych", 2)
    p(doc,
      "Model danych systemu vpnVPN zdefiniowany jest za pomocą Prisma ORM i obejmuje "
      "12 tabel w bazie PostgreSQL. Kluczowe encje to:")

    p(doc,
      "- User — użytkownik systemu z powiązaniem do konta Stripe (stripeCustomerId),")
    p(doc,
      "- Subscription — subskrypcja Stripe z informacjami o planie (tier) i statusie,")
    p(doc,
      "- Device — urządzenie klienckie z kluczem publicznym i przypisanym serwerem,")
    p(doc,
      "- VpnServer — węzeł VPN z danymi o endpointach, kluczu publicznym i statusie,")
    p(doc,
      "- VpnPeer — peer sieciowy łączący użytkownika z serwerem (klucz publiczny, dozwolone IP),")
    p(doc,
      "- VpnToken — token rejestracyjny dla węzłów VPN ze śledzeniem użycia,")
    p(doc,
      "- VpnMetric — metryka czasowa (CPU, pamięć, aktywne sesje) dla danego serwera.")

    add_figure_placeholder(doc, "Diagram relacji encji (ERD) systemu vpnVPN", "2")

    # Listing schematu Prisma (kluczowe modele)
    prisma_ranges = [(14, 32), (76, 89), (105, 123), (127, 188)]
    add_code(
        doc,
        read_code_excerpt("packages/db/prisma/schema.prisma", prisma_ranges),
        "Kluczowe modele bazy danych w schemacie Prisma",
        "3.1",
        source_path="packages/db/prisma/schema.prisma",
        line_ranges=prisma_ranges,
    )

    # 3.5 Kontrakty API
    add_heading(doc, "3.5. Kontrakty API", 2)
    p(doc,
      "Komunikacja między komponentami odbywa się za pośrednictwem API REST (Control Plane) "
      "i tRPC (Frontend). Poniższa tabela przedstawia endpointy Control Plane.")

    add_table(doc,
        ["Endpoint", "Metoda", "Uwierzytelnianie", "Opis"],
        [
            ["POST /server/register", "POST", "Bearer token", "Rejestracja węzła VPN"],
            ["GET /server/peers", "GET", "Bearer token", "Pobranie peerów dla serwera"],
            ["GET /servers", "GET", "API key", "Lista wszystkich serwerów"],
            ["DELETE /servers/:id", "DELETE", "API key", "Usunięcie serwera"],
            ["POST /peers", "POST", "API key", "Utworzenie/aktualizacja peera"],
            ["POST /peers/revoke-for-user", "POST", "API key", "Odwołanie peerów użytkownika"],
            ["DELETE /peers/:publicKey", "DELETE", "API key", "Odwołanie konkretnego peera"],
            ["GET /tokens", "GET", "API key", "Lista tokenów"],
            ["POST /tokens", "POST", "API key", "Utworzenie tokenu"],
            ["DELETE /tokens/:token", "DELETE", "API key", "Odwołanie tokenu"],
            ["POST /metrics/vpn", "POST", "Brak", "Przyjęcie metryk VPN"],
        ],
        "Endpointy Control Plane API", "4")

    p(doc,
      "Aplikacja webowa komunikuje się z Control Plane za pośrednictwem modułu "
      "controlPlane.ts, a z użytkownikiem — przez procedury tRPC. Główne routery tRPC to:")

    add_table(doc,
        ["Router", "Procedury", "Typ autoryzacji"],
        [
            ["device", "list, register, confirmConnection, cancelConnection, revoke", "paidProcedure"],
            ["billing", "createCheckoutSession, createPortalSession", "protectedProcedure"],
            ["account", "get, updateProfile, updateNotifications", "protectedProcedure"],
            ["servers", "list", "paidProcedure"],
            ["proxies", "list", "paidProcedure"],
            ["admin", "listServers, deleteServer, listTokens, createToken, revokeToken", "adminProcedure"],
            ["analytics", "summary, historicalMetrics, geoDistribution, serverHealth, recentActivity", "adminProcedure"],
            ["desktop", "resolveCode", "publicProcedure"],
        ],
        "Procedury tRPC aplikacji webowej", "5")

    # 3.6 Porównanie technologii
    add_heading(doc, "3.6. Porównanie wybranych technologii", 2)
    p(doc,
      "W trakcie projektowania systemu vpnVPN przeanalizowano i porównano różne "
      "technologie i narzędzia pod kątem wymagań projektu. Poniższe tabele prezentują "
      "kluczowe porównania, które wpłynęły na decyzje projektowe.")

    add_heading(doc, "3.6.1. Języki programowania serwera VPN", 3)

    add_table(doc,
        ["Język", "Zalety w kontekście VPN", "Wady", "Użycie w projekcie"],
        [
            ["Rust", "Bezpieczeństwo pamięci, brak GC, wydajność C/C++",
             "Stroma krzywa uczenia", "Główny język serwera VPN"],
            ["Go", "Prosty model współbieżności, bogaty ekosystem sieciowy",
             "GC, mniejsza kontrola niskopoziomowa", "Rozważany jako alternatywa"],
            ["C", "Maksymalna kontrola, najwyższa wydajność",
             "Brak bezpieczeństwa pamięci, podatność na błędy", "Tylko w komponentach zewnętrznych"],
        ],
        "Porównanie języków implementacji serwera VPN", "6")

    add_heading(doc, "3.6.2. Frameworky warstwy webowej", 3)

    add_table(doc,
        ["Framework", "Zalety", "Wady", "Użycie"],
        [
            ["Next.js (App Router)", "SSR/ISR, integracja z Vercel, SEO, full-stack",
             "Dodatkowa złożoność vs SPA", "Wybrany"],
            ["Remix", "Silny REST/HTTP focus, nowoczesny routing",
             "Mniejszy ekosystem", "Analizowany"],
            ["SPA (React + Vite)", "Prostota, elastyczność",
             "Wymaga osobnego backendu, gorsze SEO", "Użyty w desktop"],
        ],
        "Porównanie frameworków frontendowych", "7")

    add_heading(doc, "3.6.3. Platformy wdrożeniowe", 3)

    add_table(doc,
        ["Platforma", "Komponent", "Zalety", "Wady"],
        [
            ["Vercel", "Frontend (Next.js)", "Automatyczne HTTPS, CDN, edge functions",
             "Ograniczenia serverless"],
            ["Railway", "Control Plane", "Prosty deploy Docker, auto-deploy z GitHub",
             "Mniejsza kontrola niż AWS"],
            ["Manualne VM", "VPN Server", "Pełna kontrola, NET_ADMIN, host networking",
             "Wymaga ręcznego provisioningu"],
            ["GitHub Releases", "Desktop Client", "Multi-platform builds, CI/CD",
             "Brak auto-update (wymaga implementacji)"],
        ],
        "Platformy wdrożeniowe komponentów vpnVPN", "8")

    # 3.7 Diagramy sekwencji
    add_heading(doc, "3.7. Przepływy danych", 2)
    p(doc,
      "Poniższe diagramy sekwencji przedstawiają kluczowe przepływy danych w systemie "
      "vpnVPN.")

    add_heading(doc, "3.7.1. Przepływ subskrypcji", 3)
    p(doc,
      "Proces zakupu subskrypcji: (1) Użytkownik wybiera plan na stronie z cennikiem. "
      "(2) Frontend tworzy sesję Stripe Checkout przez tRPC. (3) Użytkownik dokonuje "
      "płatności na stronie Stripe. (4) Stripe wysyła webhook checkout.session.completed "
      "do backendu. (5) Backend aktualizuje subskrypcję w bazie danych. (6) Użytkownik "
      "otrzymuje e-mail potwierdzający.")

    add_figure_placeholder(doc, "Diagram sekwencji — przepływ subskrypcji", "3")

    add_heading(doc, "3.7.2. Rejestracja urządzenia i peera", 3)
    p(doc,
      "Dodanie urządzenia VPN: (1) Użytkownik klika 'Dodaj urządzenie' w dashboardzie. "
      "(2) Dla klienta webowego serwer generuje parę kluczy WireGuard, natomiast aplikacja "
      "desktopowa może przekazać klucz publiczny wygenerowany lokalnie. "
      "(3) Rekord Device jest zapisywany w bazie danych. "
      "(4) Control Plane jest informowany o nowym peerze (POST /peers). "
      "(5) Konfiguracja VPN (.conf) jest generowana i udostępniana do pobrania.")

    add_figure_placeholder(doc, "Diagram sekwencji — provisioning peera", "4")

    add_heading(doc, "3.7.3. Rejestracja węzła VPN i synchronizacja", 3)
    p(doc,
      "Cykl życia węzła VPN: (1) Przy starcie węzeł uruchamia backendy VPN i generuje "
      "klucze. (2) Węzeł rejestruje się w Control Plane (POST /server/register). "
      "(3) Co 2 sekundy węzeł pobiera listę peerów (GET /server/peers). (4) Peery "
      "są aplikowane do wszystkich aktywnych backendów (WireGuard: wg syncconf lub pełna "
      "rekonfiguracja, OpenVPN: aktualizacja secrets.txt, IKEv2: aktualizacja ipsec.secrets "
      "i swanctl --load-all). (5) Co 2 minuty "
      "węzeł wysyła heartbeat (ponowna rejestracja). (6) Control Plane oznacza serwery "
      "jako offline po 5 minutach bez heartbeatu.")

    add_figure_placeholder(doc, "Diagram sekwencji — rejestracja węzła i synchronizacja peerów", "5")


def chapter_4(doc: Document):
    """Rozdział 4: Implementacja"""
    add_heading(doc, "4. Implementacja platformy vpnVPN", 1)

    p(doc,
      "W niniejszym rozdziale przedstawiono szczegóły implementacji poszczególnych "
      "komponentów systemu vpnVPN z fragmentami kodu źródłowego. Wszystkie listingi "
      "pochodzą z rzeczywistego kodu produkcyjnego projektu.")

    # -----------------------------------------------------------------------
    # 4.1 Frontend
    # -----------------------------------------------------------------------
    add_heading(doc, "4.1. Aplikacja webowa (Next.js)", 2)
    p(doc,
      "Aplikacja webowa oparta jest na Next.js 15 z App Router i wykorzystuje TypeScript, "
      "Tailwind CSS, tRPC v11 i Prisma ORM. Wdrożona jest na platformie Vercel z "
      "automatycznym wdrażaniem z gałęzi main (produkcja) i staging (środowisko testowe).")

    # 4.1.1 Uwierzytelnianie
    add_heading(doc, "4.1.1. System uwierzytelniania", 3)
    p(doc,
      "System uwierzytelniania oparty jest o NextAuth.js z trzema dostawcami: GitHub, "
      "Google (OAuth 2.0) oraz e-mail (magic link przez Resend). Dodatkowo "
      "zaimplementowano specjalny przepływ logowania dla aplikacji desktopowej, "
      "wykorzystujący 6-cyfrowy kod OTP.")

    p_indent(doc,
      "Warstwa autoryzacji tRPC definiuje cztery poziomy procedur z rosnącymi "
      "wymaganiami dostępu: publicProcedure (bez autoryzacji), protectedProcedure "
      "(zalogowany użytkownik), paidProcedure (aktywna subskrypcja lub admin) oraz "
      "adminProcedure (rola admin).")

    trpc_init_ranges = [(67, 153)]
    add_code(
        doc,
        read_code_excerpt("apps/web/lib/trpc/init.ts", trpc_init_ranges),
        "Warstwa autoryzacji tRPC (init.ts)",
        "4.1",
        source_path="apps/web/lib/trpc/init.ts",
        line_ranges=trpc_init_ranges,
    )

    p(doc,
      "Procedura paidProcedure implementuje hierarchię dostępu: administratorzy "
      "otrzymują automatycznie uprawnienia poziomu enterprise, użytkownicy bez subskrypcji "
      "otrzymują dostęp do planu free (1 urządzenie), a pozostali — zgodnie z ich "
      "aktywnym planem subskrypcyjnym.")

    # 4.1.2 Stripe
    add_heading(doc, "4.1.2. Integracja z systemem płatności Stripe", 3)
    p(doc,
      "Obsługa płatności realizowana jest przez Stripe z mechanizmem Checkout "
      "(tworzenie sesji płatności), Customer Portal (zarządzanie subskrypcją) "
      "i webhooków (automatyczna synchronizacja statusu).")

    billing_ranges = [(1, 76)]
    add_code(
        doc,
        read_code_excerpt("apps/web/lib/trpc/routers/billing.ts", billing_ranges),
        "Tworzenie sesji Stripe Checkout (billing.ts)",
        "4.2",
        source_path="apps/web/lib/trpc/routers/billing.ts",
        line_ranges=billing_ranges,
    )

    p(doc,
      "Warstwa tRPC pokazana w listingu odpowiada za tworzenie sesji Checkout i portalu "
      "klienta. Synchronizacja statusu subskrypcji odbywa się osobno w handlerze "
      "webhooków Stripe w pliku apps/web/app/api/webhooks/stripe/route.ts, który "
      "obsługuje m.in. zdarzenia checkout.session.completed, "
      "customer.subscription.updated i customer.subscription.deleted.")

    p_indent(doc,
      "Proces tworzenia sesji Checkout obejmuje: weryfikację konfiguracji Stripe, "
      "pobranie lub utworzenie klienta Stripe (customer) powiązanego z użytkownikiem "
      "vpnVPN, a następnie utworzenie sesji Checkout z podanym identyfikatorem ceny "
      "(priceId). Sesja konfigurowana jest z przekierowaniem na dashboard po sukcesie "
      "lub na stronę cennika po anulowaniu. Włączona jest obsługa kodów promocyjnych "
      "(allow_promotion_codes: true).")

    # 4.1.3 Konfiguracja planów
    add_heading(doc, "4.1.3. Konfiguracja planów subskrypcyjnych", 3)

    tiers_ranges = [(1, 64)]
    add_code(
        doc,
        read_code_excerpt("apps/web/lib/tiers.ts", tiers_ranges),
        "Definicje planów subskrypcyjnych (tiers.ts)",
        "4.3",
        source_path="apps/web/lib/tiers.ts",
        line_ranges=tiers_ranges,
    )

    add_table(doc,
        ["Plan", "Cena (USD/mies.)", "Limit urządzeń", "Kluczowe funkcje"],
        [
            ["Free", "0", "1", "Ograniczona przepustowość, dostęp do wszystkich serwerów"],
            ["Basic", "10", "1", "Nieograniczona przepustowość, wsparcie e-mail"],
            ["Pro", "30", "5", "5 urządzeń, priorytetowe wsparcie, zaawansowane metryki"],
            ["Enterprise", "1000", "999", "Dedykowane wsparcie, niestandardowe integracje"],
        ],
        "Plany subskrypcyjne vpnVPN", "9")

    # 4.1.4 Zrzuty ekranu
    add_heading(doc, "4.1.4. Interfejs użytkownika", 3)
    p(doc,
      "Poniższe zrzuty ekranu prezentują kluczowe widoki aplikacji webowej.")

    add_figure_placeholder(doc, "Strona główna vpnVPN z metrykami floty serwerów", "6")
    add_figure_placeholder(doc, "Strona z cennikiem — wybór planu subskrypcyjnego", "7")
    add_figure_placeholder(doc, "Panel użytkownika — dashboard z przeglądem konta", "8")
    add_figure_placeholder(doc, "Panel użytkownika — zarządzanie urządzeniami VPN", "9")
    add_figure_placeholder(doc, "Panel użytkownika — lista serwerów VPN z metrykami", "10")
    add_figure_placeholder(doc, "Panel administratora — zarządzanie flotą serwerów", "11")

    # -----------------------------------------------------------------------
    # 4.2 Control Plane
    # -----------------------------------------------------------------------
    add_heading(doc, "4.2. Warstwa sterująca (Bun/Fastify)", 2)
    p(doc,
      "Control Plane to centralny komponent sterujący systemu, zbudowany na Bun "
      "(szybki runtime JavaScript/TypeScript) i frameworku Fastify. Cały serwis "
      "zaimplementowany jest w jednym pliku server.ts (~576 linii) i wdrożony na "
      "platformie Railway jako kontener Docker.")

    # 4.2.1 Walidacja
    add_heading(doc, "4.2.1. Walidacja danych wejściowych (Zod)", 3)
    p(doc,
      "Wszystkie dane wejściowe API są walidowane za pomocą biblioteki Zod, która "
      "zapewnia spójną walidację typów w czasie wykonania.")

    zod_ranges = [(22, 44)]
    add_code(
        doc,
        read_code_excerpt("services/control-plane/src/server.ts", zod_ranges),
        "Schematy walidacji Zod (server.ts)",
        "4.4",
        source_path="services/control-plane/src/server.ts",
        line_ranges=zod_ranges,
    )

    # 4.2.2 Rejestracja serwera
    add_heading(doc, "4.2.2. Rejestracja węzła VPN", 3)
    p(doc,
      "Endpoint POST /server/register umożliwia węzłom VPN samoobsługową rejestrację "
      "w systemie. Węzeł przesyła swój identyfikator, klucz publiczny WireGuard, port "
      "nasłuchiwania oraz metadane (region, kraj, adresy endpointów dla poszczególnych "
      "protokołów). Control Plane dokonuje operacji upsert na rekordzie serwera "
      "i zwiększa licznik użycia tokenu.")

    register_ranges = [(71, 83), (159, 229)]
    add_code(
        doc,
        read_code_excerpt("services/control-plane/src/server.ts", register_ranges),
        "Kontrakt i obsługa endpointu POST /server/register (server.ts)",
        "4.5",
        source_path="services/control-plane/src/server.ts",
        line_ranges=register_ranges,
    )

    # 4.2.3 Synchronizacja peerów
    add_heading(doc, "4.2.3. Synchronizacja peerów", 3)
    p(doc,
      "Endpoint GET /server/peers zwraca listę aktywnych peerów przypisanych do danego "
      "serwera. Węzły VPN odpytują ten endpoint co 2 sekundy, a Control Plane zwraca "
      "peery w formacie snake_case oczekiwanym przez serwer Rust. Peery filtrowane są "
      "po serverId lub są oznaczone jako 'unpinned' (serverId=null).")

    # -----------------------------------------------------------------------
    # 4.3 VPN Server (Rust)
    # -----------------------------------------------------------------------
    add_heading(doc, "4.3. Serwer VPN (Rust)", 2)
    p(doc,
      "Serwer VPN to program napisany w języku Rust, wykorzystujący runtime asynchroniczny "
      "Tokio i framework HTTP Axum. Kod źródłowy liczy około 3557 linii w 25 plikach "
      "i obsługuje trzy protokoły VPN przez zunifikowany interfejs cechy (trait) VpnBackend.")

    # 4.3.1 CLI
    add_heading(doc, "4.3.1. Interfejs wiersza poleceń (CLI)", 3)
    p(doc,
      "Konfiguracja serwera VPN odbywa się za pomocą argumentów wiersza poleceń (biblioteka "
      "Clap) oraz zmiennych środowiskowych. Program obsługuje dwa podpolecenia: run "
      "(uruchomienie serwera) i doctor (diagnostyka zależności systemowych).")

    cli_ranges = [(17, 61)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/main.rs", cli_ranges),
        "Konfiguracja CLI serwera VPN (main.rs) — Clap",
        "4.6",
        source_path="apps/vpn-server/src/main.rs",
        line_ranges=cli_ranges,
    )

    # 4.3.2 VpnBackend
    add_heading(doc, "4.3.2. Cecha VpnBackend i struktura PeerSpec", 3)
    p(doc,
      "Kluczowym elementem architektury serwera VPN jest cecha (trait) VpnBackend, "
      "definiująca zunifikowany interfejs dla wszystkich backendów protokołów VPN. "
      "Każdy backend implementuje metody start(), stop(), status(), apply_peers() "
      "i opcjonalnie public_key().")

    backend_ranges = [(29, 60)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/vpn/mod.rs", backend_ranges),
        "Cecha VpnBackend i struktura PeerSpec (vpn/mod.rs)",
        "4.7",
        source_path="apps/vpn-server/src/vpn/mod.rs",
        line_ranges=backend_ranges,
    )

    p(doc,
      "Struktura PeerSpec reprezentuje peera pobranego z Control Plane i zawiera "
      "klucz publiczny WireGuard, opcjonalny klucz wstępnie współdzielony (PSK), "
      "dozwolone adresy IP, oraz login i hasło dla OpenVPN/IKEv2. Ta wspólna struktura "
      "pozwala na jednolite przekazywanie konfiguracji do wszystkich backendów.")

    # 4.3.3 WireGuard
    add_heading(doc, "4.3.3. Backend WireGuard", 3)
    p(doc,
      "Backend WireGuard zarządza interfejsem wg0 i konfiguruje go za pomocą narzędzi "
      "systemowego menedżera interfejsów oraz polecenia wg syncconf (aktualizacja peerów "
      "bez restartu). "
      "Konfiguracja obejmuje: pulę adresów 10.8.0.1/24 (IPv4) i fd42:42:42::1/64 (IPv6), "
      "MTU 1420 oraz PersistentKeepalive 25 sekund.")

    wireguard_ranges = [(255, 343)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/vpn/wireguard.rs", wireguard_ranges),
        "Zastosowanie peerów WireGuard z użyciem wg syncconf (wireguard.rs)",
        "4.8",
        source_path="apps/vpn-server/src/vpn/wireguard.rs",
        line_ranges=wireguard_ranges,
    )

    p(doc,
      "Metoda apply_peers generuje kompletną konfigurację WireGuard (sekcje [Interface] "
      "i [Peer]) do pliku tymczasowego, ustawia restrykcyjne uprawnienia (0600), "
      "a następnie używa polecenia wg syncconf do atomowej aktualizacji konfiguracji "
      "interfejsu bez konieczności jego restartu. W przypadku niepowodzenia syncconf "
      "następuje powrót do pełnej rekonfiguracji przez warstwę zarządzającą interfejsem.")

    # 4.3.4 OpenVPN
    add_heading(doc, "4.3.4. Backend OpenVPN", 3)
    p(doc,
      "Backend OpenVPN konfiguruje serwer w trybie server mode z uwierzytelnianiem "
      "przez login/hasło. Hasła przechowywane są w formacie sól:hash (SHA-256) w pliku "
      "/etc/openvpn/secrets.txt, a weryfikacja odbywa się przez zewnętrzny skrypt "
      "verify.sh wywoływany przez mechanizm auth-user-pass-verify.")

    openvpn_ranges = [(103, 147)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/vpn/openvpn.rs", openvpn_ranges),
        "Skrypt weryfikacji haseł i funkcja haszowania (openvpn.rs)",
        "4.9",
        source_path="apps/vpn-server/src/vpn/openvpn.rs",
        line_ranges=openvpn_ranges,
    )

    p(doc,
      "Mechanizm haszowania wykorzystuje 16-bajtowy losowy sól (salt) dla każdego "
      "użytkownika, co zapobiega atakom słownikowym (rainbow table). Skrypt weryfikujący "
      "otrzymuje plik tymczasowy z loginem i hasłem od OpenVPN, odczytuje odpowiedni "
      "wpis z secrets.txt, oblicza hash SHA-256 z solą i porównuje z zapisanym. "
      "Sam plik secrets.txt otrzymuje w aktualnej implementacji uprawnienia 0604, "
      "ponieważ skrypt verify.sh działa z uprawnieniami użytkownika nobody i musi mieć "
      "możliwość odczytu.")

    # 4.3.5 IKEv2
    add_heading(doc, "4.3.5. Backend IKEv2/IPsec (strongSwan)", 3)
    p(doc,
      "Obsługa IKEv2 w repozytorium jest rozdzielona między dwa moduły: apps/vpn-server/src/ikev2.rs "
      "generuje bazową konfigurację strongSwan i uruchamia usługę, a backend "
      "apps/vpn-server/src/vpn/ipsec.rs aktualizuje plik /etc/ipsec.secrets z danymi "
      "użytkowników. Podstawowy profil klienta korzysta z EAP-MSCHAPv2 (login i hasło).")

    ipsec_ranges = [(88, 134)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/vpn/ipsec.rs", ipsec_ranges),
        "Zarządzanie sekretami IKEv2/strongSwan (ipsec.rs)",
        "4.10",
        source_path="apps/vpn-server/src/vpn/ipsec.rs",
        line_ranges=ipsec_ranges,
    )

    p(doc,
      "Metoda apply_peers zachowuje istniejącą linię z kluczem prywatnym serwera "
      "(RSA lub ECDSA) i dodaje wpisy EAP dla każdego peera w formacie "
      "'użytkownik : EAP \"hasło\"'. Po zapisaniu pliku z uprawnieniami 0600 "
      "wywoływane jest polecenie swanctl --load-all, które przeładowuje konfigurację "
      "sekretów bez ponownego generowania całego pliku ipsec.conf.")

    # 4.3.6 PKI
    add_heading(doc, "4.3.6. Infrastruktura klucza publicznego (PKI)", 3)
    p(doc,
      "System vpnVPN generuje własny urząd certyfikacji (CA) i certyfikat serwera "
      "przy pierwszym uruchomieniu węzła VPN. Certyfikaty wykorzystywane są przez "
      "OpenVPN (TLS) i IKEv2 (uwierzytelnianie serwera).")

    pki_ranges = [(43, 130)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/pki.rs", pki_ranges),
        "Generowanie PKI — ECDSA P-256 (pki.rs)",
        "4.11",
        source_path="apps/vpn-server/src/pki.rs",
        line_ranges=pki_ranges,
    )

    p(doc,
      "Funkcja ensure_pki generuje certyfikat CA i certyfikat serwera przy użyciu "
      "algorytmu ECDSA P-256 (PKCS_ECDSA_P256_SHA256). Certyfikat serwera zawiera "
      "publiczny adres IP jako Subject Alternative Name (SAN), co jest wymagane "
      "przez klientów IKEv2. Klucz prywatny CA nie jest zapisywany na dysku — "
      "używany jest wyłącznie do podpisania certyfikatu serwera, po czym jest usuwany "
      "z pamięci. Klucz prywatny serwera opakowywany jest w typ Zeroizing, który "
      "automatycznie zeruje pamięć przy dealokacji.")

    # 4.3.7 NAT
    add_heading(doc, "4.3.7. Konfiguracja NAT i przekazywania pakietów", 3)
    p(doc,
      "Aby klienci VPN mogli korzystać z Internetu, serwer musi skonfigurować "
      "przekazywanie pakietów (IP forwarding) i translację adresów (NAT).")

    nat_ranges = [(100, 160), (167, 255)]
    add_code(
        doc,
        read_code_excerpt("apps/vpn-server/src/main.rs", nat_ranges),
        "Konfiguracja NAT i przekazywania IP (main.rs)",
        "4.12",
        source_path="apps/vpn-server/src/main.rs",
        line_ranges=nat_ranges,
    )

    p(doc,
      "Funkcja setup_nat_and_forwarding: (1) włącza przekazywanie IPv4 i IPv6 "
      "przez sysctl; (2) wykrywa domyślny interfejs sieciowy hosta; (3) waliduje "
      "nazwę interfejsu (zapobieganie iniekcji argumentów iptables); (4) konfiguruje "
      "maskaradę NAT (iptables -t nat -A POSTROUTING -o <iface> -j MASQUERADE); "
      "(5) zezwala na ruch VPN→Internet i powrotny w łańcuchu FORWARD.")

    # -----------------------------------------------------------------------
    # 4.4 Desktop Client
    # -----------------------------------------------------------------------
    add_heading(doc, "4.4. Aplikacja desktopowa (Tauri)", 2)
    p(doc,
      "Aplikacja desktopowa vpnVPN oparta jest na frameworku Tauri 2.9, który łączy "
      "natywny webview systemu operacyjnego z backendem Rust. W przeciwieństwie do "
      "Electron, Tauri nie zawiera silnika przeglądarki, co skutkuje znacząco mniejszym "
      "rozmiarem instalatora (~5 MB vs ~150 MB).")

    add_heading(doc, "4.4.1. Architektura GUI i daemona", 3)
    p(doc,
      "Kluczową decyzją architektoniczną jest podział na dwa procesy: (1) nieuprzywilejowany "
      "interfejs graficzny (Tauri GUI) renderujący widoki React/Vite, oraz "
      "(2) uprzywilejowany daemon Rust uruchomiony z prawami root/admin, zarządzający "
      "połączeniami VPN na poziomie systemowym.")

    p(doc,
      "Daemon nasłuchuje na gnieździe Unix (/var/run/vpnvpn-daemon.sock w produkcji, "
      "/tmp/vpnvpn-daemon.sock w trybie deweloperskim) i obsługuje żądania JSON-RPC "
      "od GUI. Komunikacja obejmuje polecenia takie jak: connect (nawiązanie połączenia VPN), "
      "disconnect (rozłączenie), get_status (stan usługi) oraz get_connection_status "
      "(stan bieżącego tunelu).")

    add_figure_placeholder(doc, "Aplikacja desktopowa vpnVPN — ekran połączenia", "12")

    add_heading(doc, "4.4.2. Wsparcie wieloplatformowe", 3)
    p(doc,
      "Aplikacja desktopowa budowana jest na trzy platformy: macOS (DMG, x64 i ARM), "
      "Linux (deb, rpm, AppImage) i Windows (MSI, NSIS). Buildy realizowane są "
      "automatycznie przez GitHub Actions z macierzą platformową. Dystrybucja odbywa "
      "się przez GitHub Releases z linkami do pobrania na stronie głównej vpnVPN.")


def chapter_5(doc: Document):
    """Rozdział 5: Integracja, testowanie i wdrożenie"""
    add_heading(doc, "5. Integracja, testowanie i wdrożenie", 1)

    # 5.1 Lokalne środowisko
    add_heading(doc, "5.1. Lokalne środowisko deweloperskie", 2)
    p(doc,
      "Rozwój systemu vpnVPN odbywa się w lokalnym środowisku opartym o Docker Compose, "
      "które uruchamia pełny stos aplikacji: bazę danych PostgreSQL, Control Plane, "
      "aplikację webową, węzeł VPN oraz testowego klienta WireGuard. Wszystkie komponenty "
      "komunikują się przez rzeczywiste API HTTP — nie stosuje się mocków.")

    compose_ranges = [(5, 38), (47, 90), (92, 126), (129, 176)]
    add_code(
        doc,
        read_code_excerpt("local/compose.yaml", compose_ranges),
        "Konfiguracja Docker Compose — lokalne środowisko (compose.yaml)",
        "5.1",
        source_path="local/compose.yaml",
        line_ranges=compose_ranges,
    )

    p(doc,
      "Węzeł VPN uruchamiany jest w trybie host network (network_mode: host), co "
      "daje mu dostęp do interfejsu sieciowego hosta — niezbędne do prawidłowego "
      "działania NAT. Kontener wymaga uprawnień NET_ADMIN i dostępu do /dev/net/tun.")

    # 5.2 Testy E2E
    add_heading(doc, "5.2. Testy end-to-end", 2)
    p(doc,
      "Testy E2E weryfikują pełną ścieżkę od uruchomienia infrastruktury do nawiązania "
      "połączenia VPN. Infrastruktura testowa zdefiniowana jest w osobnym pliku "
      "compose-e2e.yaml i obejmuje scenariusze dla wszystkich trzech protokołów VPN.")

    p(doc, "Testowane scenariusze obejmują:")
    p(doc, "- Połączenie WireGuard — uruchomienie klienta z wg-quick, ping do serwera (10.8.0.1),")
    p(doc, "- Połączenie OpenVPN — uwierzytelnienie login/hasło, weryfikacja tunelu,")
    p(doc, "- Połączenie IKEv2 — uwierzytelnienie EAP-MSCHAPv2, weryfikacja tunelu,")
    p(doc, "- Odwołanie peera — weryfikacja, że po odwołaniu peera ruch VPN jest blokowany,")
    p(doc, "- Ponowne połączenie — weryfikacja odporności na przerwy w łączności.")

    p_indent(doc,
      "Testy uruchamiane są automatycznie w ramach pipeline'u CI/CD (GitHub Actions) "
      "przy każdym wysłaniu zmian lub otwarciu pull requestu dla głównych gałęzi repozytorium. "
      "Środowisko testowe jest w pełni izolowane "
      "dzięki dedykowanej sieci Docker.")

    p_indent(doc,
      "Scenariusz IKEv2 jest uruchamiany w tym samym pipeline'ie, jednak w obecnej "
      "wersji skryptu może zakończyć się statusem SKIPPED, jeśli klient strongSwan "
      "odrzuci samopodpisany certyfikat serwera w środowisku kontenerowym. Oznacza to, "
      "że test ten potwierdza gotowość ścieżki integracyjnej, ale nie jest jeszcze "
      "tak stabilny jak scenariusze WireGuard i OpenVPN.")

    p_indent(doc,
      "Kluczową decyzją projektową było zastosowanie rzeczywistych komponentów w testach "
      "zamiast mocków. Klient testowy (vpn-test-client) jest pełnoprawnym kontenerem "
      "z zainstalowanym WireGuard, który zestawia tunel VPN i weryfikuje łączność "
      "przez ping do serwera (10.8.0.1). Takie podejście pozwala wykryć problemy "
      "konfiguracyjne, sieciowe i kryptograficzne, które byłyby niewidoczne "
      "w testach opartych na mockach.")

    p_indent(doc,
      "Dla testów OpenVPN i IKEv2 stosowane są dedykowane kontenery klienckie "
      "z odpowiednimi narzędziami (openvpn, strongswan-swanctl). Każdy scenariusz "
      "testowy obejmuje: (1) oczekiwanie na gotowość serwera VPN (polling /health), "
      "(2) pobranie konfiguracji (klucz publiczny, certyfikaty), (3) zestawienie tunelu, "
      "(4) weryfikację łączności, (5) opcjonalnie — test odwołania peera i weryfikację "
      "utraty łączności.")

    # 5.3 CI/CD
    add_heading(doc, "5.3. Ciągła integracja i wdrażanie (CI/CD)", 2)
    p(doc,
      "System vpnVPN wykorzystuje GitHub Actions jako platformę CI/CD z następującymi "
      "pipeline'ami:")

    add_table(doc,
        ["Pipeline", "Wyzwalacz", "Opis"],
        [
            ["ci.yml", "Push/PR do main/staging", "Linting (Biome), testy jednostkowe, build"],
            ["deploy-backend.yml", "Push do main/staging dla apps/vpn-server/** lub workflow_dispatch", "Build obrazu Rust, push do GHCR"],
            ["desktop-build.yml", "Push do main/staging dla apps/desktop/** lub workflow_dispatch", "Macierzowy build (macOS/Linux/Windows)"],
            ["e2e.yml", "Push/PR do main/staging", "Testy E2E w Docker Compose"],
        ],
        "Pipeline'y CI/CD (GitHub Actions)", "10")

    p_indent(doc,
      "Pipeline ci.yml jest wyzwalany przy każdym push i pull request. Wykonuje "
      "kolejno: (1) instalację zależności (bun install), (2) linting kodu (Biome — "
      "alternatywa dla ESLint i Prettier, zapewniająca jednolity styl kodu), "
      "(3) testy jednostkowe (Vitest dla TypeScript, cargo test dla Rust), (4) build "
      "wszystkich pakietów (Turborepo — równoległy build z cache). Pipeline musi "
      "zakończyć się sukcesem przed możliwością mergowania pull requesta.")

    p_indent(doc,
      "Pipeline deploy-backend.yml odpowiada za budowę obrazu Docker serwera VPN "
      "(wieloetapowy build: Rust 1.88 builder + Debian bookworm-slim runtime) i "
      "publikację do GitHub Container Registry (GHCR). Obraz zawiera binarki "
      "WireGuard-tools, OpenVPN i strongSwan. Pipeline desktop-build.yml buduje "
      "aplikację desktopową na macierzy trzech platform (macOS, Linux, Windows) "
      "i publikuje artefakty jako GitHub Release.")

    # 5.4 Wdrożenie
    add_heading(doc, "5.4. Wdrożenie produkcyjne", 2)
    p(doc,
      "Wdrożenie systemu vpnVPN obejmuje cztery niezależne procesy wdrożenia:")

    add_heading(doc, "5.4.1. Aplikacja webowa i panel administracyjny", 3)
    p(doc,
      "Aplikacja Next.js wdrażana jest na platformie Vercel z automatycznym wdrożeniem "
      "z gałęzi main (produkcja: vpnvpn.dev) i staging (staging.vpnvpn.dev). Vercel "
      "zapewnia CDN, automatyczne certyfikaty HTTPS, edge functions i środowiska podglądowe "
      "dla każdego pull requesta.")

    add_heading(doc, "5.4.2. Warstwa sterująca", 3)
    p(doc,
      "Control Plane wdrożony jest na platformie Railway jako kontener Docker z "
      "automatycznym wdrożeniem z GitHub. Konfiguracja zapisana jest w pliku railway.toml, "
      "który definiuje proces budowy, health check i port nasłuchiwania. Domena "
      "produkcyjna: api.vpnvpn.dev.")

    add_heading(doc, "5.4.3. Węzły VPN", 3)
    p(doc,
      "Węzły VPN wdrażane są ręcznie na maszynach wirtualnych za pomocą skryptu "
      "setup-vpn-node.sh, który: (1) instaluje Docker, (2) pobiera obraz z GHCR "
      "(ghcr.io/xnetcat/vpnvpn/vpn-server), (3) konfiguruje zmienne środowiskowe "
      "(API_URL, VPN_TOKEN, SERVER_ID), (4) uruchamia kontener z host networking "
      "i uprawnieniami NET_ADMIN. Po uruchomieniu węzeł automatycznie rejestruje się "
      "w Control Plane.")

    add_heading(doc, "5.4.4. Aplikacja desktopowa", 3)
    p(doc,
      "Buildy aplikacji desktopowej generowane są automatycznie przez GitHub Actions "
      "z macierzą platformową (macOS x64/ARM, Linux x64, Windows x64) i publikowane "
      "jako GitHub Releases. Linki do pobrania prezentowane są na stronie głównej "
      "vpnVPN z automatyczną detekcją platformy użytkownika.")

    add_figure_placeholder(doc, "Metryki Prometheus eksportowane przez serwer VPN", "13")


def chapter_6(doc: Document):
    """Rozdział 6: Podsumowanie"""
    add_heading(doc, "6. Podsumowanie", 1)

    add_heading(doc, "6.1. Ocena realizacji założeń", 2)
    p(doc,
      "W ramach niniejszej pracy zaprojektowano i zaimplementowano kompletną platformę "
      "vpnVPN — system SaaS do zarządzania usługą VPN. Platforma składa się z czterech "
      "głównych komponentów: aplikacji webowej (Next.js), warstwy sterującej (Bun/Fastify "
      "z PostgreSQL), serwera VPN (Rust) obsługującego trzy protokoły (WireGuard, OpenVPN, "
      "IKEv2/IPsec) oraz aplikacji desktopowej (Tauri z daemonem Rust).")

    p_indent(doc,
      "Cele zdefiniowane we wstępie zostały w zdecydowanej większości osiągnięte. "
      "Repozytorium zawiera działający panel webowy z autoryzacją i rozliczeniami, "
      "warstwę sterującą z kontraktami API, węzeł VPN obsługujący trzy protokoły oraz "
      "aplikację desktopową komunikującą się z lokalnym daemonem.")

    p_indent(doc,
      "Ocena wymagań niefunkcjonalnych wymaga bardziej ostrożnego wniosku. System nie "
      "loguje treści ruchu sieciowego i stosuje ochronę materiału kryptograficznego "
      "(m.in. ograniczenia uprawnień plików i typ Zeroizing), ale dla protokołów "
      "OpenVPN i IKEv2 przechowuje również dane uwierzytelniające potrzebne do "
      "zestawienia tunelu. Architektura typu pull upraszcza skalowanie węzłów, choć "
      "nie eliminuje całkowicie pracy operacyjnej po stronie administratora.")

    add_table(doc,
        ["Cel szczegółowy", "Ocena realizacji", "Potwierdzenie w repozytorium"],
        [
            ["Opracowanie architektury platformy SaaS dla VPN", "Zrealizowany",
             "Rozdziały 3 i 4 oraz moduły apps/web, services/control-plane, apps/vpn-server"],
            ["Implementacja centralnego zarządzania użytkownikami i subskrypcjami", "Zrealizowany",
             "NextAuth.js, Stripe, router billing, dashboard i modele Prisma"],
            ["Implementacja wieloprotokołowego węzła VPN", "Zrealizowany",
             "Backendy WireGuard, OpenVPN i IKEv2 oraz wspólny model PeerSpec"],
            ["Przygotowanie klienta desktopowego dla użytkownika końcowego", "Zrealizowany częściowo",
             "Aplikacja Tauri i daemon działają, lecz część funkcji pozostaje rozwojowa"],
            ["Weryfikacja rozwiązania w środowisku lokalnym i CI", "Zrealizowany częściowo",
             "Docker Compose, testy E2E i workflowy GitHub Actions; scenariusz IKEv2 bywa pomijany"],
        ],
        "Ocena realizacji celów pracy", "11")

    add_heading(doc, "6.2. Ograniczenia", 2)
    p(doc,
      "Zidentyfikowane ograniczenia systemu obejmują:")

    p(doc,
      "- Ręczne wdrażanie węzłów VPN — proces provisioningu wymaga ręcznego "
      "uruchomienia skryptu na maszynie wirtualnej. Automatyzacja tego procesu "
      "(np. integracja z Terraform/Pulumi) jest zaplanowana, ale nie została "
      "zrealizowana w ramach pracy.")

    p(doc,
      "- Brak klientów mobilnych — system posiada klientów web i desktop, ale "
      "nie oferuje dedykowanej aplikacji mobilnej (iOS/Android). Użytkownicy mobilni "
      "mogą korzystać z natywnego IKEv2 lub klientów WireGuard/OpenVPN.")

    p(doc,
      "- Zależność od zewnętrznych usług — platforma zależy od Vercel (frontend), "
      "Railway (Control Plane), Stripe (płatności) i Neon (baza danych), co "
      "wprowadza ryzyko vendor lock-in.")

    p(doc,
      "- Przechowywanie danych uwierzytelniających dla części protokołów — "
      "obsługa OpenVPN i IKEv2 wymaga w aktualnej architekturze przechowywania "
      "nazw użytkowników i haseł, co zwiększa wymagania dotyczące ochrony danych.")

    add_heading(doc, "6.3. Kierunki dalszego rozwoju", 2)
    p(doc,
      "Zidentyfikowane kierunki dalszego rozwoju obejmują:")

    p(doc,
      "- Klienci mobilni — natywne aplikacje iOS i Android z obsługą WireGuard "
      "i IKEv2, zintegrowane z systemem subskrypcji.")

    p(doc,
      "- Automatyczne wdrażanie węzłów — integracja z Terraform lub Pulumi "
      "umożliwiająca dodawanie nowych węzłów VPN bezpośrednio z panelu administratora.")

    p(doc,
      "- Split tunneling — selektywne kierowanie ruchu przez VPN (np. tylko "
      "określone domeny lub aplikacje).")

    p(doc,
      "- Wsparcie proxy SOCKS5/HTTP — tunelowanie ruchu TCP przez proxy, "
      "umożliwiające obejście blokad DPI.")

    p(doc,
      "- Publiczny dashboard metryk — weryfikowalny, publicznie dostępny "
      "dashboard prezentujący zagregowane metryki systemu, zwiększający "
      "transparentność usługi.")

    p(doc,
      "- Auto-update aplikacji desktopowej — mechanizm automatycznej aktualizacji "
      "w oparciu o Tauri Updater i GitHub Releases.")


def add_bibliography(doc: Document):
    """Add bibliography."""
    add_heading(doc, "Bibliografia", 1)

    refs = [
        '[1] J. A. Donenfeld, "WireGuard: Next Generation Kernel Network Tunnel", '
        'Proceedings of the 2017 Network and Distributed System Security Symposium (NDSS), 2017.',

        '[2] RFC 7296 \u2014 C. Kaufman, P. Hoffman, Y. Nir, P. Eronen, T. Kivinen, '
        '"Internet Key Exchange Protocol Version 2 (IKEv2)", Internet Engineering Task Force, 2014.',

        '[3] RFC 4555 \u2014 P. Eronen, "IKEv2 Mobility and Multihoming Protocol (MOBIKE)", '
        'Internet Engineering Task Force, 2006.',

        '[4] RFC 7748 \u2014 A. Langley, M. Hamburg, S. Turner, "Elliptic Curves for '
        'Security", Internet Engineering Task Force, 2016.',

        '[5] RFC 8439 \u2014 Y. Nir, A. Langley, "ChaCha20 and Poly1305 for IETF Protocols", '
        'Internet Engineering Task Force, 2018.',

        '[6] RFC 7693 \u2014 J.-P. Aumasson, S. Neves, Z. Wilcox-O\'Hearn, '
        '"The BLAKE2 Cryptographic Hash and Message Authentication Code (MAC)", '
        'Internet Engineering Task Force, 2015.',

        '[7] OpenVPN Inc., "OpenVPN 2.6 Reference Manual", '
        'https://openvpn.net/community-resources/reference-manual-for-openvpn-2-6/, '
        'dost\u0119p: 26.03.2026.',

        '[8] strongSwan Project, "strongSwan Documentation", '
        'https://docs.strongswan.org/, dost\u0119p: 26.03.2026.',

        '[9] Next.js, "Next.js Documentation", Vercel Inc., '
        'https://nextjs.org/docs, dost\u0119p: 26.03.2026.',

        '[10] Tauri Contributors, "Tauri Documentation", https://tauri.app/start/, '
        'dost\u0119p: 26.03.2026.',

        '[11] S. Klabnik, C. Nichols, "The Rust Programming Language", '
        'https://doc.rust-lang.org/book/, dost\u0119p: 26.03.2026.',

        '[12] Tokio Contributors, "Tokio Documentation", https://tokio.rs/, '
        'dost\u0119p: 26.03.2026.',

        '[13] Prisma, "Prisma ORM Documentation", https://www.prisma.io/docs, '
        'dost\u0119p: 26.03.2026.',

        '[14] Stripe, "Subscriptions", https://stripe.com/docs/billing/subscriptions, '
        'dost\u0119p: 26.03.2026.',

        '[15] T. Perrin, "The Noise Protocol Framework", https://noiseprotocol.org/noise.html, '
        'dost\u0119p: 26.03.2026.',

        '[16] Rozporz\u0105dzenie Parlamentu Europejskiego i Rady (UE) 2016/679 z dnia '
        '27 kwietnia 2016 r. w sprawie ochrony os\u00f3b fizycznych w zwi\u0105zku z przetwarzaniem '
        'danych osobowych i w sprawie swobodnego przep\u0142ywu takich danych.',

        '[17] Fastify Contributors, "Fastify Documentation", https://fastify.dev/docs/, '
        'dost\u0119p: 26.03.2026.',

        '[18] Bun, "Bun Documentation", https://bun.sh/docs, dost\u0119p: 26.03.2026.',

        '[19] Railway, "Railway Docs", https://docs.railway.com/, dost\u0119p: 26.03.2026.',

        '[20] Vercel, "Vercel Documentation", https://vercel.com/docs, '
        'dost\u0119p: 26.03.2026.',
    ]

    for ref in refs:
        doc.add_paragraph(ref, style=STYLE_BIB)


# ===========================================================================
# MAIN
# ===========================================================================

def main():
    print("Generowanie pracy licencjackiej...")

    # Create main body document
    doc = Document()
    setup_styles(doc)
    setup_page(doc)
    add_page_numbers(doc)
    enable_update_fields_on_open(doc)

    # Table of contents
    add_toc(doc)

    # Chapters
    chapter_1(doc)
    doc.add_page_break()
    chapter_2(doc)
    doc.add_page_break()
    chapter_3(doc)
    doc.add_page_break()
    chapter_4(doc)
    doc.add_page_break()
    chapter_5(doc)
    doc.add_page_break()
    chapter_6(doc)
    doc.add_page_break()
    add_bibliography(doc)

    # Save body separately so it can be appended after the original title page
    doc.save(str(BODY_FILE))
    print(f"  Zapisano treść roboczą: {BODY_FILE}")

    if TITLE_PAGE.exists():
        print(f"  Scalanie ze stroną tytułową: {TITLE_PAGE}")
        title_doc = Document(str(TITLE_PAGE))
        strip_toc_bookmarks(title_doc)
        enable_update_fields_on_open(title_doc)
        body_section = title_doc.add_section(WD_SECTION.NEW_PAGE)
        setup_page(title_doc, section=body_section)
        add_page_numbers(title_doc, section=body_section)
        body_doc = Document(str(BODY_FILE))
        composer = Composer(title_doc)
        composer.append(body_doc)
        composer.save(str(OUTPUT_FILE))
        print(f"  Zapisano kompletny dokument: {OUTPUT_FILE}")
    else:
        doc.save(str(OUTPUT_FILE))
        print(f"  UWAGA: Brak strony tytułowej ({TITLE_PAGE}), zapisano samą treść.")

    final_doc = Document(str(OUTPUT_FILE))
    enable_update_fields_on_open(final_doc)
    final_doc.save(str(OUTPUT_FILE))

    if BODY_FILE.exists():
        BODY_FILE.unlink()

    print("Gotowe!")


if __name__ == "__main__":
    main()
